# -*- coding: utf-8 -*-
"""
Markdown转Word转换器核心类
"""

import re
import logging
from pathlib import Path
from typing import List, Tuple, Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .config import ConverterConfig


# ==================== 格式化工具函数 ====================

def set_font(run, font_name: str, size: Pt = None, bold: bool = None):
    """设置中文字体"""
    run.font.name = font_name

    # 设置颜色为黑色
    r = run._element
    if r.rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    else:
        rPr = r.rPr

    color_element = OxmlElement('w:color')
    color_element.set(qn('w:val'), '000000')
    rPr.append(color_element)

    # 禁用下划线
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'none')
    rPr.append(underline)

    # 设置字体
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold


def set_paragraph_spacing(paragraph, space_before=None, space_after=None, line_spacing=None):
    """设置段落格式"""
    pPr = paragraph._element.get_or_add_pPr()

    if space_before or space_after or line_spacing:
        spacing = OxmlElement('w:spacing')

        if space_before:
            spacing.set(qn('w:before'), str(int(space_before)))
        if space_after:
            spacing.set(qn('w:after'), str(int(space_after)))
        if line_spacing:
            if isinstance(line_spacing, Pt):
                # 固定值：Pt转缇，1pt = 20缇
                spacing.set(qn('w:line'), str(int(line_spacing.pt * 20)))
                spacing.set(qn('w:lineRule'), 'exact')
            elif isinstance(line_spacing, (int, float)):
                # 倍数行距
                spacing.set(qn('w:line'), str(int(line_spacing * 240)))
                spacing.set(qn('w:lineRule'), 'auto')

        pPr.append(spacing)


def set_run_format(run, font_config, bold=False, italic=False):
    """设置运行格式（字体、粗体、斜体）

    Args:
        run: Word运行对象
        font_config: 字体配置
        bold: 是否加粗
        italic: 是否斜体
    """
    set_font(run, font_config.name, font_config.size, bold if bold else font_config.bold)
    if italic:
        run.font.italic = True


def set_first_line_indent_chars(paragraph, chars: int):
    """设置首行缩进（字符单位）

    Args:
        paragraph: 段落对象
        chars: 缩进字符数
    """
    pPr = paragraph._element.get_or_add_pPr()
    indent = OxmlElement('w:ind')

    # Word中字符单位：1字符 = 100 twips
    indent.set(qn('w:firstLineChars'), str(chars * 100))
    indent.set(qn('w:firstLine'), str(chars * 100))

    pPr.append(indent)


def set_page_margins(section):
    """设置页边距"""
    margins = ConverterConfig.MARGINS
    section.top_margin = Pt(margins.top * 28.35)
    section.bottom_margin = Pt(margins.bottom * 28.35)
    section.left_margin = Pt(margins.left * 28.35)
    section.right_margin = Pt(margins.right * 28.35)
    section.header_distance = Pt(margins.header * 28.35)
    section.footer_distance = Pt(margins.footer * 28.35)


def remove_paragraph_border(paragraph):
    """移除段落边框"""
    try:
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            pPr.remove(pBdr)
    except Exception:
        pass


def is_signature_line(text: str) -> bool:
    """判断是否是落款或日期行"""
    text_stripped = text.strip()
    return '辅导员：' in text_stripped or '日期：' in text_stripped


# ==================== 解析工具函数 ====================

def parse_inline_formatting(text: str) -> str:
    """清理行内格式标记,移除所有Markdown格式符号

    处理内容:
    - 移除粗体标记 **text**
    - 移除斜体标记 *text*
    - 移除反引号 `code`

    Args:
        text: 原始文本

    Returns:
        str: 清理后的纯文本
    """
    result = text
    i = 0

    while i < len(result):
        # 处理粗体 **text**
        if result[i:i+2] == '**':
            j = result.find('**', i + 2)
            if j != -1:
                # 移除 ** 和 **,保留中间内容
                result = result[:i] + result[i+2:j] + result[j+2:]
                continue
            else:
                # 没有结束标记,移除开始的 **
                result = result[:i] + result[i+2:]
                continue
        # 处理斜体 *text*
        elif result[i] == '*':
            j = result.find('*', i + 1)
            if j != -1:
                # 移除 * 和 *,保留中间内容
                result = result[:i] + result[i+1:j] + result[j+1:]
                continue
            else:
                # 没有结束标记,移除开始的 *
                result = result[:i] + result[i+1:]
                continue
        # 处理反引号 `code`
        elif result[i] == '`':
            j = result.find('`', i + 1)
            if j != -1:
                # 移除 ` 和 `,保留中间内容
                result = result[:i] + result[i+1:j] + result[j+1:]
                continue
            else:
                # 没有结束标记,移除开始的 `
                result = result[:i] + result[i+1:]
                continue
        i += 1

    return result


def extract_title_from_md(md_file_path: str) -> str:
    """从Markdown文件中提取第一个一级标题"""
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line.startswith('#'):
                    title = line.lstrip('#').strip()
                    title = re.sub(r'[<>:"/\\|?*]', '', title)
                    return title if title else '未命名文档'
    except Exception:
        pass

    return '未命名文档'


def sanitize_filename(filename: str) -> str:
    """清理文件名，移除非法字符"""
    filename = re.sub(r'[<>:"/\\|?*]', '', filename)
    filename = filename.strip()
    if len(filename) > 200:
        filename = filename[:200]
    return filename if filename else '未命名文档'


def normalize_list_symbol(text: str) -> str:
    """移除列表符号前缀

    Args:
        text: 原始文本

    Returns:
        str: 处理后的文本（移除-、+、*、•等列表符号）
    """
    # 移除常见的列表符号前缀（包括空格）
    if text.startswith(('- ', '+ ', '* ')):
        return text[2:]  # 移除符号和空格
    elif text.startswith('• '):
        return text[2:]  # 移除bullet和空格（•是1个字符，加上空格共2个）
    return text


# ==================== 转换器类 ====================


class MarkdownConverter:
    """Markdown转Word转换器"""

    def __init__(self, config: Optional[ConverterConfig] = None, logger: Optional[logging.Logger] = None):
        """
        初始化转换器

        Args:
            config: 配置对象，默认使用ConverterConfig
            logger: 日志记录器
        """
        self.config = config or ConverterConfig()
        self.logger = logger or logging.getLogger(__name__)

    def convert(self, md_file: Path) -> Tuple[Optional[Document], str]:
        """
        转换Markdown文件为Word文档

        Args:
            md_file: Markdown文件路径

        Returns:
            Tuple[Document, str]: (Word文档对象, 标题)
        """
        try:
            self.logger.info(f"开始处理: {md_file.name}")

            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()

            doc = Document()
            title = extract_title_from_md(md_file)
            self.logger.info(f"文档标题: {title}")

            # 设置页面格式
            self._setup_page_layout(doc)

            # 解析并转换内容（标题在内容中处理）
            self._parse_content(doc, content)

            return doc, title

        except Exception as e:
            self.logger.error(f"处理文件 {md_file.name} 时出错: {e}")
            return None, "未命名文档"

    def _setup_page_layout(self, doc: Document):
        """设置页面布局"""
        for section in doc.sections:
            set_page_margins(section)

    def _parse_content(self, doc: Document, content: str):
        """解析Markdown内容并添加到文档"""
        lines = content.split('\n')
        in_code_block = False
        i = 0

        while i < len(lines):
            line = lines[i].rstrip()

            # 跳过代码块
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                i += 1
                continue
            if in_code_block:
                i += 1
                continue

            # 跳过空行
            if not line:
                i += 1
                continue

            # 检测表格（以|开头且包含|）
            if self._is_table_line(line):
                # 收集整个表格
                table_lines = []
                while i < len(lines) and self._is_table_line(lines[i].rstrip()):
                    table_lines.append(lines[i].rstrip())
                    i += 1
                self._process_table(doc, table_lines)
                continue

            # 处理各种Markdown元素
            if line.startswith('# ') and not line.startswith('## '):
                self._process_heading1(doc, line)
            elif line.startswith('## ') and not line.startswith('### '):
                self._process_heading2(doc, line)
            elif line.startswith('### '):
                self._process_heading3(doc, line)
            elif line.strip().startswith('>'):
                self._process_quote(doc, line)
            elif line.strip().startswith('---') or line.strip().startswith('***'):
                self._process_separator(doc)
            elif self._is_list_item(line):
                self._process_list_item(doc, line)
            else:
                self._process_paragraph(doc, line)

            i += 1

    def _is_table_line(self, line: str) -> bool:
        """判断是否是表格行"""
        return '|' in line and (line.startswith('|') or line.strip().startswith('|'))

    def _process_table(self, doc: Document, table_lines: List[str]):
        """处理Markdown表格"""
        if not table_lines:
            return

        # 过滤掉分隔行并解析数据
        table_data = []
        for line in table_lines:
            if '---' in line:
                continue
            # 移除首尾的|，然后按|分割
            line = line.strip()
            if line.startswith('|'):
                line = line[1:]
            if line.endswith('|'):
                line = line[:-1]
            cells = [cell.strip() for cell in line.split('|')]
            if cells:
                table_data.append(cells)

        if not table_data:
            return

        # 创建Word表格
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        font_config = self.config.get_font('body')

        # 填充表格数据
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    para = cell.paragraphs[0]
                    para.clear()

                    # 清理Markdown格式标记
                    cell_text = parse_inline_formatting(cell_text)
                    run = para.add_run(cell_text)
                    set_run_format(run, font_config)

                    set_paragraph_spacing(para, line_spacing=self.config.PARAGRAPH.line_spacing)
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def _process_heading1(self, doc: Document, line: str):
        """处理一级标题"""
        heading_text = line[2:].strip()
        heading = doc.add_heading(heading_text, level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        font_config = self.config.get_font('heading1')
        for run in heading.runs:
            set_font(run, font_config.name, font_config.size, font_config.bold)

        set_paragraph_spacing(
            heading,
            space_before=self.config.PARAGRAPH.space_before_title,
            space_after=self.config.PARAGRAPH.space_after_title,
            line_spacing=self.config.PARAGRAPH.line_spacing
        )

        # 移除标题边框和小黑点
        remove_paragraph_border(heading)

    def _process_heading2(self, doc: Document, line: str):
        """处理二级标题"""
        heading_text = line[3:].strip()
        heading = doc.add_heading(heading_text, level=2)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        font_config = self.config.get_font('heading2')
        for run in heading.runs:
            set_font(run, font_config.name, font_config.size, font_config.bold)

        set_paragraph_spacing(
            heading,
            space_before=self.config.PARAGRAPH.space_before_title,
            space_after=self.config.PARAGRAPH.space_after_title,
            line_spacing=self.config.PARAGRAPH.line_spacing
        )

        # 移除标题边框和小黑点
        remove_paragraph_border(heading)

    def _process_heading3(self, doc: Document, line: str):
        """处理三级标题"""
        heading_text = line[4:].strip()
        heading = doc.add_heading(heading_text, level=3)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        font_config = self.config.get_font('heading3')
        for run in heading.runs:
            set_font(run, font_config.name, font_config.size, font_config.bold)

        set_paragraph_spacing(
            heading,
            space_before=self.config.PARAGRAPH.space_before_title,
            space_after=self.config.PARAGRAPH.space_after_title,
            line_spacing=self.config.PARAGRAPH.line_spacing
        )

        # 移除标题边框和小黑点
        remove_paragraph_border(heading)

    def _process_quote(self, doc: Document, line: str):
        """处理引用块"""
        quote_text = line.strip()[1:].strip()
        # 清理Markdown格式标记
        quote_text = parse_inline_formatting(quote_text)
        p = doc.add_paragraph(quote_text)

        set_paragraph_spacing(p, line_spacing=self.config.PARAGRAPH.line_spacing)
        set_first_line_indent_chars(p, self.config.PARAGRAPH.first_line_indent_chars)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        font_config = self.config.get_font('quote')
        for run in p.runs:
            set_run_format(run, font_config)

    def _process_separator(self, doc: Document):
        """处理分隔线"""
        p = doc.add_paragraph('_' * 50)
        set_paragraph_spacing(p, line_spacing=self.config.PARAGRAPH.line_spacing)

    def _is_list_item(self, line: str) -> bool:
        """判断是否是列表项"""
        return (line.strip().startswith(('* ', '- ', '• ', '+ ')) or
                re.match(r'^\s*\d+[\.\)]\s', line))

    def _process_list_item(self, doc: Document, line: str):
        """处理列表项"""
        list_content = normalize_list_symbol(line.strip())
        # 清理Markdown格式标记
        list_content = parse_inline_formatting(list_content)
        p = doc.add_paragraph(list_content)

        set_paragraph_spacing(p, line_spacing=self.config.PARAGRAPH.line_spacing)
        set_first_line_indent_chars(p, self.config.PARAGRAPH.first_line_indent_chars)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # 统一设置字体格式
        font_config = self.config.get_font('body')
        for run in p.runs:
            set_run_format(run, font_config)

    def _process_paragraph(self, doc: Document, line: str):
        """处理普通段落"""
        # 去除首尾空格
        line = line.strip()
        # 清理Markdown格式标记
        line = parse_inline_formatting(line)

        p = doc.add_paragraph(line)
        set_paragraph_spacing(p, line_spacing=self.config.PARAGRAPH.line_spacing)

        # 检查是否是落款或日期
        if is_signature_line(line):
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            # 落款不缩进
            pPr = p._element.get_or_add_pPr()
            indent = OxmlElement('w:ind')
            indent.set(qn('w:firstLine'), '0')
            pPr.append(indent)
        else:
            set_first_line_indent_chars(p, self.config.PARAGRAPH.first_line_indent_chars)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # 统一设置字体格式
        font_config = self.config.get_font('body')
        for run in p.runs:
            set_run_format(run, font_config)
